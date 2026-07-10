# -*- coding: utf-8 -*-
"""
Build a Chinese course report (.docx) from the RS-Token v0.5 manuscript.

Format requirements:
- Title: 三号黑体, 居中
- Body: 小四宋体, 1.5 倍行距
- Heading 1: 四号黑体
- References: GB/T 7714
- Figures and tables numbered with captions
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------- 字体助手 ----------
FONT_HEI = "黑体"
FONT_SONG = "宋体"
FONT_TIMES = "Times New Roman"

PT_SAN = Pt(16)         # 三号 = 16 pt
PT_SI = Pt(14)          # 四号 = 14 pt
PT_XIAOSI = Pt(12)      # 小四 = 12 pt
PT_WUHAO = Pt(10.5)     # 五号 = 10.5 pt
PT_XIAOWUHAO = Pt(9)    # 小五 = 9 pt

def set_run_font(run, eastasia=FONT_SONG, ascii_font=FONT_TIMES, size=PT_XIAOSI, bold=False):
    run.font.name = ascii_font
    run.font.size = size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), eastasia)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)

def add_paragraph(doc, text, *, eastasia=FONT_SONG, ascii_font=FONT_TIMES,
                  size=PT_XIAOSI, bold=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing=1.5,
                  first_line_indent_chars=2,
                  space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = space_before
    fmt.space_after = space_after
    if first_line_indent_chars > 0:
        # 中文首行缩进 2 字符
        fmt.first_line_indent = Pt(size.pt * first_line_indent_chars)
    if text:
        run = p.add_run(text)
        set_run_font(run, eastasia=eastasia, ascii_font=ascii_font, size=size, bold=bold)
    return p

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, eastasia=FONT_HEI, ascii_font=FONT_HEI, size=PT_SAN, bold=True)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, eastasia=FONT_HEI, ascii_font=FONT_HEI, size=PT_SI, bold=True)
    return p

def add_h2(doc, text):
    # 二级标题：小四黑体
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    fmt.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, eastasia=FONT_HEI, ascii_font=FONT_HEI, size=PT_XIAOSI, bold=True)
    return p

def add_caption(doc, text, *, before=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, eastasia=FONT_SONG, ascii_font=FONT_TIMES,
                 size=PT_WUHAO, bold=True)
    return p

def add_image(doc, path, *, width_cm=12.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p

def set_cell_text(cell, text, *, bold=False, size=PT_WUHAO,
                  eastasia=FONT_SONG, ascii_font=FONT_TIMES,
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, eastasia=eastasia, ascii_font=ascii_font,
                 size=size, bold=bold)

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_table(doc, header, rows, *, col_widths_cm=None, header_bold=True):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # header
    for j, txt in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, txt, bold=header_bold)
        set_cell_border(cell)
    # body
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            set_cell_text(cell, txt, bold=False)
            set_cell_border(cell)
    return table

def add_reference_paragraph(doc, idx, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Pt(0)
    # 悬挂缩进：大约 2 字符
    fmt.left_indent = Cm(0.74)
    # python-docx hanging indent
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:hanging'), '420')  # twentieths of a point ~ 21 pt = 0.74 cm
    ind.set(qn('w:left'), '420')
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(f"[{idx}] " + text)
    set_run_font(run, eastasia=FONT_SONG, ascii_font=FONT_TIMES,
                 size=PT_WUHAO, bold=False)
    return p

# ---------- 文档构建 ----------
def build():
    doc = Document()

    # 全局页面（A4，标准页边距）
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 设置默认 Normal 样式，避免后续段落继承默认 Calibri
    style = doc.styles['Normal']
    style.font.name = FONT_TIMES
    style.font.size = PT_XIAOSI
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT_SONG)
    rfonts.set(qn('w:ascii'), FONT_TIMES)
    rfonts.set(qn('w:hAnsi'), FONT_TIMES)

    # 标题
    add_title(doc, "面向信道鲁棒遥感通信的层次化语义 Token 表示研究")
    add_title.__doc__  # noqa

    # 作者 / 课程信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("课程：遥感解译       学院：电子信息工程学院       姓名：张宝辉       学号：y125211005")
    set_run_font(run, eastasia=FONT_SONG, ascii_font=FONT_TIMES, size=PT_XIAOSI, bold=False)

    # 摘要
    add_h1(doc, "摘  要")
    add_paragraph(doc,
        "遥感图像下行链路常面临带宽受限、低信噪比和衰落信道，而接收端在不同任务阶段对信息的需求差异很大：实时告警只需场景级判别，"
        "现场研判需要可大致解读的图像，离线测绘和归档则需要高保真重建。本文围绕《遥感解译》课程涉及的遥感场景理解任务，"
        "系统梳理一种面向遥感通信的层次化离散表示——RS-Token。该方法在四阶段残差向量量化（RVQ）的基础上，引入仅施加于第一层的"
        "RemoteCLIP 蒸馏约束：第 0 层（L0）承载地物语义，第 1—3 层（L1—L3）承载用于重建细化的残差信息；接收端按信道质量"
        "选择前 k 层 token 作为前缀传输。在 AID 数据集 30 类场景分类任务上，蒸馏使无信道下 L0 准确率从 58.23% 提升到 83.33%，"
        "在 AWGN +5 dB 下保持 82.57%，在 Rayleigh +10 dB 下保持 78.80%；重建 PSNR 随 k 从 1 增至 4，由 22.98 dB 提升到 25.92 dB，"
        "重建图分类准确率由 71.4% 提升到 86.8%。蒸馏位置反事实实验进一步表明，将 RemoteCLIP 信号从 L0 移至 L1 会让 L0 的"
        "可分性塌陷 48 个百分点，从而以因果方式确认了“仅蒸馏 L0”的层特化是设计选择带来的而非架构偶然。在与 rate-1/2 LDPC "
        "保护的 JPEG2000/WebP 流水线进行严格相同传输比特预算对比时，RS-Token 解码失败率为 0.0%，传统流水线则在 60—100% 的"
        "图像上失败。整套部署侧编解码器仅 10.87 M 参数，单图 GPU 推理 8.59 ms，符合无人机和边缘端的工程约束。本文从遥感解译"
        "的视角整理上述方法与实验，讨论其在低空巡检、应急筛查与归档测绘等典型场景中的解译意义。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "关键词：遥感解译；语义通信；残差向量量化；RemoteCLIP；任务导向通信；信道鲁棒性",
        bold=False, first_line_indent_chars=2)

    # 1 引言
    add_h1(doc, "1  引言")
    add_paragraph(doc,
        "遥感解译是从遥感图像中提取地物类别、目标位置和场景属性等信息的过程，是国土调查、环境监测、灾害评估和军事侦察等领域的"
        "基础环节[1][2]。随着低空无人机、卫星星座以及移动观测平台的普及，遥感图像逐渐从“事后处理的存储数据”演变为"
        "“在轨产生、按需传输”的实时数据流。这一变化对图像通信链路提出了新的要求：在带宽受限、信噪比波动和衰落明显的环境下，"
        "下行链路需要兼顾任务即时性与重建保真。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "传统“先压缩、再传输、再解译”的流程对遥感解译并不完全友好。一方面，JPEG2000、WebP 等通用压缩码流以视觉重建为优化目标，"
        "并不显式保留场景级语义[3][4]；另一方面，这些码流对未纠正的 bit error 极其敏感，无保护传输时即便是个位数的误码也可能"
        "造成整张图像解码失败[3]。即使引入信道编码，也只能降低误码率而不能改变源表示的语义优先级[5]。深度联合源信道编码"
        "（Deep JSCC）尝试以连续表示直接对抗信道噪声[6]，但其输出难以与现代数字通信系统中的离散接口、重传机制和差错保护机制"
        "无缝对接。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "另一条思路是直接传输离散神经 token 索引。VQ-VAE[7] 与残差向量量化（RVQ）[8] 自然地把图像编码成多层 codebook 索引，且"
        "RVQ 天然支持按层传输[9]。但是，标准 RVQ tokenizer 由重建损失驱动，第一层 codebook 主要承担连续 latent 的粗糙逼近，"
        "并不必然对场景类别具有判别性。如果直接把 RVQ 的 L0 索引拿去做遥感场景识别，效果通常不理想。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "本文围绕这一问题，对一种名为 RS-Token 的层次化离散表示进行系统介绍与实验复现。其核心思想是：在四层 RVQ 的基础上，"
        "仅在第一层（L0）量化输出上加入 RemoteCLIP[10] 蒸馏约束，使 L0 对齐到面向遥感图文配对训练的语义嵌入空间，从而把"
        "“地物语义”集中到最先传输的 2 560 比特上；其余三层（L1—L3）保持重建残差细化的角色。在《遥感解译》课程框架下，"
        "本文不仅复述方法本身，也尝试从遥感解译应用的角度对实验现象进行解读，讨论 L0 token 在场景级解译中的代理作用、不同"
        "信道下的解译性能下界，以及该表示对解译任务部署的潜在价值。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "本文的具体研究内容与组织结构如下：第 2 节梳理与本文密切相关的三条技术线索——任务导向语义通信、离散视觉 token 与残差量化、"
        "遥感基础模型与基准数据集；第 3 节给出问题设定、层次化 RVQ 编解码器、L0 蒸馏机制与信道模型的形式化描述；第 4 节按"
        "“蒸馏对 L0 解译保真度的影响、蒸馏位置反事实、渐进式重建与 codec+LDPC 对比、跨数据集 zero-shot 迁移、teacher 与 λ "
        "消融”五个层次组织实验；第 5 节从遥感解译应用的角度对实验结果进行讨论，尝试给出该方法在低空巡检、应急筛查与归档测绘"
        "等典型解译场景中的部署建议；第 6 节总结全文并展望后续可在变化检测、目标检测、语义分割等更细粒度解译任务上的扩展。",
        first_line_indent_chars=2)

    # 2 相关工作
    add_h1(doc, "2  相关工作")
    add_h2(doc, "2.1  任务导向遥感通信")
    add_paragraph(doc,
        "任务导向语义通信关注下游决策而非像素保真。早期 Deep JSCC[6] 工作表明端到端学习的图像编码在含噪信道下显著优于"
        "“压缩 + 经典信道码”串联流程；面向无人机遥感的语义通信[11] 进一步在更低码率下保持较高的解译任务准确率。近年研究"
        "把多模态基础模型作为语义先验，构建零样本语义通信[12]，或将多层 codebook 与生成式语义通信耦合得到数字方案[9]。"
        "本文所讨论的 RS-Token 与这条线索一致，但坚持“离散、层次化、只在第一层语义化”的设计，使其在标准数字接口下更具工程"
        "可部署性。",
        first_line_indent_chars=2)

    add_h2(doc, "2.2  离散视觉 Token 与残差量化")
    add_paragraph(doc,
        "VQ-VAE[7] 提出了离散 latent 图像 token 的概念；残差向量量化[8] 通过多层 codebook 累积逼近同一 latent，使表示"
        "天然具有“前缀可截断”的层次性，被 SoundStream[13]、EnCodec 等神经音频编解码大量采用。视觉端的 BEiT v2[14] 将"
        "CLIP 视觉特征蒸馏到 VQ tokenizer 上，证明引入基础模型语义可以使离散 token 更具判别性。RS-Token 把这一思路扩展到"
        "遥感领域，并明确把蒸馏限制在“最先传输”的那一层。",
        first_line_indent_chars=2)

    add_h2(doc, "2.3  遥感基础模型与图像数据集")
    add_paragraph(doc,
        "RemoteCLIP[10] 在大规模遥感图文对上训练，提供了与遥感地物语义对齐的视觉嵌入空间，是本文蒸馏所用的 teacher。"
        "AID[15] 与 NWPU-RESISC45[16] 是被广泛采用的遥感场景识别基准，覆盖城市、农田、水体、交通枢纽等典型解译类别，"
        "本文的全部任务路径与重建路径实验均在这两个数据集上进行。",
        first_line_indent_chars=2)

    # 3 方法
    add_h1(doc, "3  方法")
    add_h2(doc, "3.1  问题设定")
    add_paragraph(doc,
        "设遥感图像 x∈ℝ^(H×W×3)，对应场景类别 y。发送端将 x 编码为 K 层离散索引 z₁,…,z_K（本文 K=4）。所有层共享相同的 16×16 "
        "空间 token 网格，每层每张图贡献 2 560 比特。接收端按信道质量请求前 k 层前缀，所占比特预算 B(k)=2 560·k，对应 "
        "k=1,2,3,4 时为 2 560、5 120、7 680、10 240 比特/图。其中 k=1 对应低码率任务模式，k≥2 用于渐进式重建。",
        first_line_indent_chars=2)

    add_h2(doc, "3.2  层次化 RVQ 编解码器")
    add_paragraph(doc,
        "编码器 E 把图像映射到连续 latent u=E(x)，残差量化器以 r₀=u 起步，第 ℓ 层选取最近邻码字 z_ℓ=argmin_j ‖r_(ℓ-1)-c_(ℓ,j)‖²，"
        "残差更新 r_ℓ=r_(ℓ-1)-q_ℓ，其中 q_ℓ=c_(ℓ,z_ℓ) 为该层量化向量。解码器 D 在收到前 k 层时按 x̂_k=D(Σ_(ℓ=1..k) q_ℓ) 重建。"
        "由于这一“前缀性质”，同一个训练好的 tokenizer 即可支持四种码率工作点，无需重训。",
        first_line_indent_chars=2)

    add_h2(doc, "3.3  仅在 L0 上的 RemoteCLIP 蒸馏")
    add_paragraph(doc,
        "设 f_T(x) 为冻结的 RemoteCLIP 图像嵌入（已 L2 归一化），g(q₁) 为从 L0 量化输出到 teacher 维度的轻量投影头，"
        "蒸馏损失定义为 L_distill = 1 − cos(g(q₁), f_T(x))。整体训练目标为 L = L_rec + L_vq + λ·L_distill，主模型 λ=0.5；"
        "无蒸馏内部基线 rvq_baseline 取 λ=0，其余完全一致。这种“仅在 L0 蒸馏”的设计是层特化的关键：它使第一层在不依赖更高层的"
        "情况下就能完成场景级解译，而其他三层不被语义任务“挤占”，仍可专注于重建细节。",
        first_line_indent_chars=2)

    add_h2(doc, "3.4  信道模型与评估指标")
    add_paragraph(doc,
        "前 k 层 RVQ 索引被串行成比特序列 b∈{0,1}^B(k)，BPSK 调制后经信道传输，y_i = h_i·s_i + n_i，AWGN 时 h_i≡1，"
        "Rayleigh 衰落时 h_i 服从 𝒞𝒩(0,1)，每符号独立采样[5]。接收端做相干硬判决恢复比特，再按 codebook 大小重组成索引。"
        "任务路径只在 k=1 下评估：对收到的 L0 索引求“词袋直方图”h₀，再训一线性分类器做 AID 30 类场景识别；重建路径则用 "
        "PSNR、LPIPS[17] 和一个在干净 AID 上独立训练的 ResNet34 重建图分类器（test top-1 96.10%）评估 k=1…4。",
        first_line_indent_chars=2)

    # 图 1
    fig1_path = r"d:\CODE\遥感+通信\遥感+通信\rstoken\figs\fig_method_aech_image2pro.png"
    if os.path.exists(fig1_path):
        add_image(doc, fig1_path, width_cm=14.0)
        add_caption(doc, "图 1  RS-Token 总体框架（来源：本文复现实验，导出自 rstoken/figs/fig_method_aech_image2pro.png）")

    # 4 实验
    add_h1(doc, "4  实验")
    add_h2(doc, "4.1  数据与训练设置")
    add_paragraph(doc,
        "实验使用 AID 30 类航空场景数据集[15]，按类分层划分 8 000/1 000/1 000 的训练/验证/测试集，每类训练 176—336 张、"
        "验证测试每类 22—42 张。输入分辨率 256×256，latent 网格 16×16，每层 codebook 大小 1 024，因此每层每图 256·log₂(1024)="
        "2 560 比特。优化器为 AdamW，学习率 1×10⁻⁴，批大小 16，训练 50 个 epoch，重建/LPIPS/RVQ/蒸馏的损失权重分别为 1、0.1、1、0.5。"
        "Teacher 为冻结的 RemoteCLIP-ViT-B/32，蒸馏投影头隐藏维度 512。所有信道实验均不插入纠错码，单独的 LDPC 对比实验在 4.4 节"
        "报告。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "为保证可比性，所有内部消融模型均共享相同的初始化种子（除显式的多种子实验外），相同的随机数据增广（训练用 RandomResizedCrop "
        "256，评估用 Resize+CenterCrop 256），相同的 bf16 自动混合精度，相同的梯度裁剪阈值。AID 与 NWPU-RESISC45 数据集均做了"
        "按类分层划分，避免不同 split 的类别分布偏置；信道仿真在 PyTorch 端实现，BPSK 调制、AWGN/Rayleigh 信道、相干硬判决均按 [5] "
        "中标准定义。蒸馏位置、teacher、λ 三组 ablation 仅以蒸馏配置不同进行单变量对比，其余训练超参保持完全一致。",
        first_line_indent_chars=2)

    # 4.2 蒸馏对 L0 任务保真度的影响
    add_h2(doc, "4.2  RemoteCLIP 蒸馏对 L0 解译保真度的影响")
    add_paragraph(doc,
        "本组实验比较仅蒸馏开关不同的 rvq_baseline 与 rvq_distill，在 41/42/43 三个独立训练种子下报告 L0 词袋准确率 mean±std。"
        "结果（表 1、图 2）显示，蒸馏后 L0 在无信道下从 58.23%±1.57% 提升到 83.33%±0.81%，绝对提升 25.10 个百分点；"
        "在 AWGN +5 dB 下仍保持 82.57%±0.31%，几乎不衰减；在 Rayleigh +10 dB 下保持 78.80%±0.72%，相较基线提升 30.17 pp。"
        "由于两组模型共享相同初始化种子，可以排除“权重初始化噪声”这一备选解释，将提升归因于蒸馏本身。",
        first_line_indent_chars=2)

    add_caption(doc, "表 1  不同信道条件下 L0 任务路径准确率（h₀，单位：%）")
    table1_header = ["指标", "rvq_baseline", "rvq_distill"]
    table1_rows = [
        ["无信道", "58.23±1.57", "83.33±0.81"],
        ["AWGN +5 dB", "55.73±0.72", "82.57±0.31"],
        ["AWGN +10 dB", "58.20±1.59", "83.37±0.76"],
        ["Rayleigh +5 dB", "28.67±0.65", "58.57±0.47"],
        ["Rayleigh +10 dB", "48.63±1.31", "78.80±0.72"],
        ["Best PSNR (dB)", "26.10±0.02", "25.89±0.07"],
        ["Best LPIPS", "0.172±0.001", "0.175±0.002"],
    ]
    add_table(doc, table1_header, table1_rows, col_widths_cm=[4.5, 4.5, 4.5])

    fig2_path = r"d:\CODE\遥感+通信\遥感+通信\rstoken\figs\fig_exp_l0_task_robustness_v3.png"
    if os.path.exists(fig2_path):
        add_image(doc, fig2_path, width_cm=14.0)
        add_caption(doc, "图 2  L0 任务路径在 AWGN/Rayleigh 信道下的鲁棒性（来源：rstoken/figs/fig_exp_l0_task_robustness_v3.png）")

    add_paragraph(doc,
        "从遥感解译角度看，这一结果意味着只用最先到达的 2 560 比特，接收端就能在 +5 dB AWGN 这一相对噪声较大的链路上实现 80% 以上的"
        "30 类场景判别准确率。这对应急救援、灾损初判、低空巡检中的“先粗判后细看”需求十分契合：值守端可以在网络刚连上、噪声仍较高时"
        "迅速确定该幅图像属于哪一类典型地物，再决定是否拉取后续层进行精细解译。",
        first_line_indent_chars=2)

    # 4.3 蒸馏位置反事实
    add_h2(doc, "4.3  蒸馏位置反事实：L0、全部层、L1 三种方案的对比")
    add_paragraph(doc,
        "为验证“L0 特化”的因果归属，本文在严格相同的种子、teacher、训练超参条件下，对比把蒸馏分别加在 L0（主模型）、所有四层 z_q "
        "之和（all layers）、仅 L1 三种方案。表 2 报告 h₀ 准确率、k=4 PSNR、k=4 重建图分类准确率以及在累积 codeword 上的层次线性 probe。",
        first_line_indent_chars=2)

    add_caption(doc, "表 2  蒸馏位置反事实实验（n=1 000 AID test/信道，单种子 42）")
    t2_header = ["指标", "L0 only（主）", "All layers", "L1 only"]
    t2_rows = [
        ["h₀ 无信道 (%)", "82.6", "81.2", "34.2"],
        ["h₀ AWGN +5 dB (%)", "82.1", "79.8", "32.1"],
        ["h₀ Rayleigh +5 dB (%)", "58.6", "49.3", "10.4"],
        ["h₀ Rayleigh +10 dB (%)", "77.7", "74.2", "24.0"],
        ["k=4 PSNR 无信道 (dB)", "25.92", "25.64", "22.32"],
        ["k=4 recon-cls 无信道 (%)", "86.9", "86.1", "35.3"],
        ["k=4 recon-cls AWGN +5 dB (%)", "84.6", "83.1", "3.7"],
        ["probe L0 (%)", "86.0", "86.2", "30.5"],
        ["probe L0+L1+L2+L3 (%)", "87.9", "88.3", "36.5"],
    ]
    add_table(doc, t2_header, t2_rows, col_widths_cm=[5.0, 3.2, 3.2, 3.2])

    add_paragraph(doc,
        "可以看到，当蒸馏被错误地施加到 L1 时，L0 的可分性彻底崩塌（无信道 h₀ 跌至 34.2%，−48.4 pp），但 L1 自己也并未补偿性地"
        "“接管语义”——L0+L1+L2+L3 累积 probe 仅 36.5%，整模型重建也同步退化。把蒸馏均匀分散到所有层只带来轻微稀释（约 −1.4 pp），"
        "尚不灾难。综合三组反事实，可以从因果意义上确认：L0 承载语义不是“恰好出现在第一层”，而是“蒸馏放在哪一层，语义就出现在哪一层”，"
        "把它放在 L0 才能与“前缀传输”这一通信目标对齐。",
        first_line_indent_chars=2)

    # 4.4 渐进式重建与传统编码基线
    add_h2(doc, "4.4  渐进式重建与同传输比特 codec+LDPC 对比")
    add_paragraph(doc,
        "重建路径上，传输更多 RVQ 层是否带来收益？表 3 报告 rvq_distill 在三种子下的 PSNR、LPIPS 与重建图分类准确率，并与"
        "rate-1/2 LDPC 保护下的 JPEG2000、WebP 进行严格相同传输比特预算的对比。",
        first_line_indent_chars=2)

    add_caption(doc, "表 3  RS-Token 重建路径与同传输比特 codec+LDPC stress baseline（节选）")
    t3_header = ["条件", "k / 方法", "Total bits", "解码失败 (%)", "PSNR (dB)", "Recon-cls (%)"]
    t3_rows = [
        ["无信道", "RS-Token k=1", "2 560", "0.0", "22.98±0.09", "71.4±1.3"],
        ["无信道", "RS-Token k=4", "10 240", "0.0", "25.92±0.08", "86.8±0.4"],
        ["AWGN +5 dB", "RS-Token k=1", "2 560", "0.0", "21.99±0.07", "68.5±1.4"],
        ["AWGN +5 dB", "RS-Token k=4", "10 240", "0.0", "23.94±0.09", "84.4±0.9"],
        ["Rayleigh +10 dB", "RS-Token k=4", "10 240", "0.0", "20.55±0.07", "67.7±1.1"],
        ["AWGN +10 dB", "RS-Token k=4", "20 480", "0.0", "25.92", "86.9"],
        ["AWGN +10 dB", "JPEG2000+LDPC", "20 480", "62.1", "—", "14.3"],
        ["AWGN +10 dB", "WebP+LDPC", "20 480", "79.4", "—", "13.6"],
        ["Rayleigh +10 dB", "RS-Token k=4", "20 480", "0.0", "25.08", "86.3"],
        ["Rayleigh +10 dB", "JPEG2000+LDPC", "20 480", "71.6", "—", "1.3"],
        ["Rayleigh +10 dB", "WebP+LDPC", "20 480", "100.0", "—", "0.0"],
    ]
    add_table(doc, t3_header, t3_rows,
              col_widths_cm=[2.6, 3.2, 2.0, 2.4, 2.6, 2.6])

    fig3_path = r"d:\CODE\遥感+通信\遥感+通信\rstoken\figs\fig_exp_progressive_reconstruction_v3.png"
    if os.path.exists(fig3_path):
        add_image(doc, fig3_path, width_cm=14.0)
        add_caption(doc, "图 3  RS-Token 在 k=1…4 下的渐进式重建（来源：rstoken/figs/fig_exp_progressive_reconstruction_v3.png）")

    add_paragraph(doc,
        "无信道下随 k 从 1 增至 4，PSNR 由 22.98 dB 升至 25.92 dB（gap 2.94 dB，远大于每个 cell 的标准差 0.04—0.09 dB），"
        "重建图分类准确率由 71.4% 升至 86.8%；AWGN +5 dB 下趋势一致；Rayleigh +5 dB 衰落较深，渐进性大体丧失，所有 k 下的 PSNR "
        "均在 16.9 dB 附近。Rayleigh +10 dB 部分恢复，k=4 的 PSNR 达到 20.55 dB。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "在与 codec+LDPC 流水线的同比特对比中，RS-Token 在所有信道与所有传输比特预算下解码失败率均为 0%；JPEG2000+LDPC 在 AWGN "
        "+10 dB、传输 20 480 比特时仍有 62.1% 的图像解码失败，WebP+LDPC 失败 79.4%；Rayleigh +10 dB 下 WebP+LDPC 全失败。这一差距"
        "并不源自纠错强度的不同——双方使用同一码——而源自表示形态：RS-Token 的离散索引可以逐 token 查表恢复，单点比特错只波及该 token 的"
        "码字选择；而 JPEG2000/WebP 的熵编码容器一旦被破坏，整个码流便无法解析。",
        first_line_indent_chars=2)

    # 4.5 跨数据集
    add_h2(doc, "4.5  跨数据集 zero-shot 迁移")
    add_paragraph(doc,
        "为检验蒸馏所得的 L0 是否真承载“通用遥感语义”，本文在 AID 上训练好的 tokenizer 不更新任何参数（编码器、RVQ codebook、解码器"
        "全部冻结），只在 NWPU-RESISC45[16] 训练集上重新拟合一线性 h₀ probe，再在 6 300 张 NWPU 测试图上评估。",
        first_line_indent_chars=2)

    add_caption(doc, "表 4  L0 tokenizer 到 NWPU-RESISC45 的 zero-shot 迁移（单种子 42）")
    t4_header = ["条件", "rvq_distill (%)", "rvq_baseline (%)", "提升 (pp)"]
    t4_rows = [
        ["无信道", "64.4", "43.5", "20.9"],
        ["AWGN +5 dB", "61.4", "41.9", "19.5"],
        ["AWGN +10 dB", "64.4", "43.5", "20.9"],
        ["Rayleigh +5 dB", "29.7", "16.5", "13.2"],
        ["Rayleigh +10 dB", "53.3", "33.6", "19.7"],
    ]
    add_table(doc, t4_header, t4_rows, col_widths_cm=[3.6, 3.4, 3.4, 2.5])

    add_paragraph(doc,
        "AID 无信道下蒸馏的 gap 是 25.1 pp，迁移到 NWPU-RESISC45 的 45 类任务后只缩小到 20.9 pp，绝对蒸馏增益保留约 83%。"
        "考虑 NWPU 类别更细（45 类 vs 30 类，机会值 2.2% vs 3.3%）且包含 AID 中不存在的若干场景，这一结果支持“蒸馏所得 L0 编码"
        "的是遥感领域的通用场景语义，而非 AID-specific 特征”。",
        first_line_indent_chars=2)

    # 4.6 Teacher / 损失权重消融
    add_h2(doc, "4.6  Teacher 选择与蒸馏权重 λ 的消融")
    add_paragraph(doc,
        "Teacher 消融对比 RemoteCLIP 与 OpenAI CLIP（结构相同、种子相同），结果（表 5）显示 RemoteCLIP 仅在 h₀ 上有 1.4—6.0 pp 的"
        "优势，且优势随衰落增大；k=4 重建 PSNR 与重建图分类基本相同，差距分别在 0.1 dB 与 0.3 pp 之内。可以认为，承担机制的是“视觉语言"
        "蒸馏”这一类 teacher，RemoteCLIP 提供的是一个领域适配的 teacher，其增益主要在更具挑战性的衰落信道下显现。",
        first_line_indent_chars=2)

    add_caption(doc, "表 5  Teacher 消融与蒸馏权重 λ sweep（节选）")
    t5_header = ["条件", "RemoteCLIP λ=0.5", "OpenAI CLIP λ=0.5", "λ=0.1", "λ=1.0"]
    t5_rows = [
        ["无信道 h₀ (%)", "82.6", "80.8", "71.2", "84.5"],
        ["AWGN +5 dB h₀ (%)", "82.5", "78.8", "69.2", "83.3"],
        ["Rayleigh +5 dB h₀ (%)", "59.8", "53.8", "42.1", "49.8"],
        ["Rayleigh +10 dB h₀ (%)", "79.2", "74.4", "64.3", "77.6"],
        ["k=4 无信道 PSNR (dB)", "25.92", "26.03", "26.20", "25.64"],
        ["k=4 无信道 recon-cls (%)", "86.9", "86.6", "—", "—"],
    ]
    add_table(doc, t5_header, t5_rows, col_widths_cm=[3.6, 2.8, 2.8, 2.0, 2.0])

    add_paragraph(doc,
        "λ sweep 表明，单调提高蒸馏权重对无信道与 AWGN 任务有益，但在 Rayleigh 衰落下 λ=1.0 反而劣于 λ=0.5，PSNR 与 LPIPS 也"
        "随 λ 增大而单调下降。这是“语义对齐”与“衰落鲁棒性”之间的折衷点，本文的 λ=0.5 即在该 trade-off 上选择中庸方案。",
        first_line_indent_chars=2)

    # 5 解译应用层面的讨论
    add_h1(doc, "5  从遥感解译角度的讨论")
    add_paragraph(doc,
        "上述实验回答了通信侧的问题，但本课程的核心是“遥感解译”。结合解译任务，本文做以下三点讨论。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "（1）L0 token 可作为解译流水线的低成本入口。在传统遥感解译中，初步分类往往在抽稀图像或低分辨率预览上完成，再对感兴趣"
        "区域进行高分辨率解译。RS-Token 的 L0 词袋特征 h₀ 在严格 2 560 比特下提供 30 类（AID）/45 类（NWPU）场景识别能力，可看作"
        "“云端轻量解译入口”：解译员或自动管线在收到 L0 后即可完成场景类别筛选、初步任务分流，再决定是否拉取 L1—L3 进行高分辨率"
        "解译，从而压缩链路负载。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "（2）信道条件直接决定解译能力的上界。表 3 显示 Rayleigh +5 dB 下即使全部传输也只能得到 16.9 dB 的 PSNR 与约 20% 的"
        "重建图分类准确率，远低于干净基线。这一结果提示我们，在低空 UAV 巡检、应急筛查这类 link budget 紧张的场景，解译产品的"
        "“可信度”需要与链路质量联动评估；表 1 中 L0 词袋在 Rayleigh +10 dB 下仍达 78.80%，说明在这些场景下应当把“场景级判别”"
        "作为可靠产品，把“像素级重建解译”作为机会型产品。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "（3）跨数据集迁移意味着该框架可作为遥感解译模型的“信道适配前置层”。NWPU-RESISC45 上仍能保留 20.9 pp 的蒸馏增益，"
        "说明 L0 token 不是只在 AID 类目集合上有效。对实际部署而言，可以把 RS-Token 视为“信道感知的语义编码层”：上层解译模型"
        "（场景分类、目标检测、语义分割）只需要在它的输出上重训轻量 head 即可适配，节约下游模型的迁移成本。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "（4）端到端部署的工程友好性。整套部署侧编解码器仅 10.87 M 参数、单图 GPU 推理 8.59 ms、CPU 单线程 66.70 ms；"
        "RemoteCLIP teacher（约 150 M 参数）只在训练阶段使用，部署阶段无需加载，这与无人机和边缘计算节点的算力/功耗约束相吻合。"
        "接收端只需保存约 8 MB 的 codebook（fp32），即可逐 token 查表恢复 L0—L3 对应的量化向量，再按需走任务路径或重建路径。"
        "对于解译业务而言，这意味着同一个轻量化模型可以同时服务“低码率告警通道”和“高码率精解译通道”，无需为不同业务训练不同模型，"
        "运维成本显著降低。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "（5）与传统遥感图像处理流程的兼容性。RS-Token 输出的是离散索引比特流，可直接接入现有的 OFDM、HARQ、自适应调制编码"
        "（AMC）等数字物理层；同时，由于 L0 已经具备语义判别能力，可以与传统的 K-Means 聚类、主题模型、视觉词袋分类器等遥感解译"
        "经典算法结合，构建“神经 token + 经典分类器”的混合解译流水线，使该方法不必绑定特定的深度学习上层模型，亦能在工程项目中"
        "渐进式落地。",
        first_line_indent_chars=2)

    # 6 结论
    add_h1(doc, "6  结论与展望")
    add_paragraph(doc,
        "本文围绕《遥感解译》课程主题，对一种面向遥感通信的层次化离散表示——RS-Token——进行了系统介绍与实验复盘。通过仅在第一层"
        "RVQ 上施加 RemoteCLIP 蒸馏，使 L0 集中承载场景级解译语义，余下三层用于渐进式重建细化。AID 上的实验证实蒸馏带来约 25 pp 的"
        "L0 解译准确率提升，且在 AWGN/Rayleigh 信道下保持稳定；蒸馏位置反事实从因果上确认 L0 特化来自“蒸馏放在哪一层”而非架构偶然；"
        "与同比特 codec+LDPC 流水线的对比则凸显离散 token 表示在含噪信道下解码鲁棒性的本质差距。NWPU-RESISC45 上的 zero-shot 迁移"
        "进一步表明蒸馏 L0 是真实的遥感语义而非 AID-specific 特征。",
        first_line_indent_chars=2)

    add_paragraph(doc,
        "局限方面，当前任务路径证据集中在场景级解译，尚未覆盖目标检测、语义分割等更细粒度解译任务；多 seed 复现仅在主重建实验"
        "上完成，蒸馏位置/teacher/跨数据集消融为单种子；LDPC 仅是自定义 rate-1/2 稀疏码而非 5G NR LDPC。下一步工作包括：把 L0 token "
        "扩展到面向变化检测的时序解译任务；联合优化层选、调制和信道编码；探索遥感多模态（光学+SAR）扩展方案。",
        first_line_indent_chars=2)

    # 参考文献
    add_h1(doc, "参考文献")
    refs = [
        # 1
        "李德仁, 童庆禧, 李荣兴, 等. 高分辨率对地观测的若干前沿科学问题[J]. 中国科学: 地球科学, 2012, 42(6): 805–813.",
        # 2
        "孙显, 王智睿, 孙元睿, 等. 遥感图像智能解译: 从单模态到多模态[J]. 测绘学报, 2024, 53(4): 643–662.",
        # 3
        "Taubman D S, Marcellin M W. JPEG2000: Image Compression Fundamentals, Standards and Practice[M]. Boston: Kluwer Academic Publishers, 2002.",
        # 4
        "Google Inc. WebP: A new image format for the Web[EB/OL]. (2023-09-01)[2026-06-01]. https://developers.google.com/speed/webp.",
        # 5
        "Proakis J G, Salehi M. Digital Communications[M]. 5th ed. New York: McGraw-Hill, 2007.",
        # 6
        "Bourtsoulatze E, Kurka D B, Gündüz D. Deep joint source-channel coding for wireless image transmission[J]. IEEE Transactions on Cognitive Communications and Networking, 2019, 5(3): 567–579.",
        # 7
        "van den Oord A, Vinyals O, Kavukcuoglu K. Neural discrete representation learning[C]//Advances in Neural Information Processing Systems (NeurIPS). Long Beach, CA, USA: Curran Associates, 2017: 6306–6315.",
        # 8
        "Lee D, Kim C, Kim S, et al. Autoregressive image generation using residual quantization[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR). New Orleans, LA, USA: IEEE, 2022: 11523–11532.",
        # 9
        "Wang T, Cheng J, Yang Y, et al. MOC-RVQ: Multi-level codebook based digital generative semantic communication[C]//Proceedings of IEEE GLOBECOM. Cape Town, South Africa: IEEE, 2024: 3215–3220.",
        # 10
        "Liu F, Chen D, Guan Z, et al. RemoteCLIP: A vision-language foundation model for remote sensing[J]. IEEE Transactions on Geoscience and Remote Sensing, 2024, 62: 5622216.",
        # 11
        "Gao Y, Liu Y, Lyu B, et al. Task-oriented image transmission for UAV scene classification in disaster response[J]. IEEE Transactions on Communications, 2022, 70(8): 5181–5196.",
        # 12
        "Yilmaz S F, Hasırcıoğlu B, Gündüz D. Zero-shot semantic communication with multimodal foundation models[J]. IEEE Transactions on Vehicular Technology, 2026, 75(5): 7280–7295.",
        # 13
        "Zeghidour N, Luebs A, Omran A, et al. SoundStream: An end-to-end neural audio codec[J]. IEEE/ACM Transactions on Audio, Speech, and Language Processing, 2022, 30: 495–507.",
        # 14
        "Peng Z, Dong L, Bao H, et al. BEiT v2: Masked image modeling with vector-quantized visual tokenizers[EB/OL]. (2022-08-12)[2026-06-01]. https://arxiv.org/abs/2208.06366.",
        # 15
        "Xia G S, Hu J, Hu F, et al. AID: A benchmark data set for performance evaluation of aerial scene classification[J]. IEEE Transactions on Geoscience and Remote Sensing, 2017, 55(7): 3965–3981.",
        # 16
        "Cheng G, Han J, Lu X. Remote sensing image scene classification: Benchmark and state of the art[J]. Proceedings of the IEEE, 2017, 105(10): 1865–1883.",
        # 17
        "Zhang R, Isola P, Efros A A, et al. The unreasonable effectiveness of deep features as a perceptual metric[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR). Salt Lake City, UT, USA: IEEE, 2018: 586–595.",
    ]
    for i, r in enumerate(refs, start=1):
        add_reference_paragraph(doc, i, r)

    # 保存
    out_path = r"d:\CODE\遥感+通信\遥感+通信\paper_draft\latex\遥感解译课程报告_RS-Token.docx"
    doc.save(out_path)
    print("OK:", out_path)
    print("size:", os.path.getsize(out_path), "bytes")

if __name__ == "__main__":
    build()
