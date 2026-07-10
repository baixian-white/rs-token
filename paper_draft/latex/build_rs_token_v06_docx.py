# -*- coding: utf-8 -*-
"""Build the formal Chinese DOCX manuscript for RS-Token v0.6."""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml


HERE = Path(__file__).resolve().parent
ROOT = HERE.parents[1]
SOURCE = HERE / "rs_token_v0.6_zh.md"
BBL = HERE / "rs_token_v06.bbl"
OUTPUT = HERE / "RS-Token_v0.6_中文稿_公式与三线表修订版.docx"
FIG_DIR = ROOT / "rstoken" / "figs"

FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_LATIN = "Times New Roman"
FONT_MATH = "Cambria Math"


def set_run_font(run, *, cn=FONT_CN, latin=FONT_LATIN, size=Pt(10.5),
                 bold=None, italic=None, color=None):
    run.font.name = latin
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


SUBSCRIPT = str.maketrans("0123456789+-=()aehijklmnoprstuvx", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
SUPERSCRIPT = str.maketrans("0123456789+-=()inT", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁱⁿᵀ")


def _script_replace(text: str, marker: str, table: dict[int, str]) -> str:
    pattern = re.compile(re.escape(marker) + r"\{([^{}]+)\}|" + re.escape(marker) + r"([A-Za-z0-9+\-=*]+)")

    def repl(match):
        value = match.group(1) or match.group(2)
        translated = value.translate(table)
        return translated if len(translated) == len(value) else marker + value

    previous = None
    while previous != text:
        previous = text
        text = pattern.sub(repl, text)
    return text


def latex_math_to_text(text: str) -> str:
    text = text.strip().strip("$")
    text = text.replace("\\begin{aligned}", "").replace("\\end{aligned}", "")
    text = text.replace("\\\\", "\n")

    def balanced_group(value: str, start: int) -> tuple[str, int] | None:
        if start >= len(value) or value[start] != "{":
            return None
        depth = 0
        for pos in range(start, len(value)):
            if value[pos] == "{":
                depth += 1
            elif value[pos] == "}":
                depth -= 1
                if depth == 0:
                    return value[start + 1:pos], pos + 1
        return None

    # Convert nested LaTeX fractions before braces are removed.
    while True:
        match = re.search(r"\\(?:t)?frac", text)
        if not match:
            break
        first = balanced_group(text, match.end())
        second = balanced_group(text, first[1]) if first else None
        if not first or not second:
            break
        numerator = latex_math_to_text(first[0])
        denominator = latex_math_to_text(second[0])
        text = text[:match.start()] + f"({numerator})/({denominator})" + text[second[1]:]

    replacements = {
        r"\mathbb{R}": "ℝ", r"\mathcal{L}": "L", r"\mathcal{CN}": "CN",
        r"\lambda": "λ", r"\ell": "ℓ", r"\beta": "β", r"\sigma": "σ",
        r"\Delta": "Δ", r"\phi": "φ", r"\psi": "ψ", r"\theta": "θ",
        r"\pm": "±", r"\times": "×", r"\in": "∈", r"\equiv": "≡",
        r"\sim": "∼", r"\top": "ᵀ", r"\to": "→", r"\ldots": "…",
        r"\cdot": "·", r"\geq": "≥", r"\leq": "≤", r"\approx": "≈",
        r"\arg\min": "arg min", r"\arg\max": "arg max", r"\sum": "Σ",
        r"\log": "log", r"\dagger": "†", r"\mathbf": "", r"\hat": "",
        r"\bar": "", r"\mathrm": "", r"\rm": "", r"\operatorname": "",
        r"\mathrm": "", r"\text": "", r"\left": "", r"\right": "",
        r"\bigl": "", r"\bigr": "", r"\Bigl": "", r"\Bigr": "",
        r"\qquad": "  ", r"\quad": "  ", r"\;": " ", r"\,": "",
        r"\!": "", r"\|": "‖", r"\{": "{", r"\}": "}", r"\%": "%",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\(?:sgn|Re|MSE|PSNR|Acc|E)\b", lambda m: m.group(0)[1:], text)
    text = re.sub(r"\\[A-Za-z]+", "", text)
    text = _script_replace(text, "_", SUBSCRIPT)
    text = _script_replace(text, "^", SUPERSCRIPT)
    text = text.replace("{", "").replace("}", "").replace("&", "")
    text = re.sub(r"[ \t]+", " ", text)
    return "\n".join(part.strip(" ,") for part in text.splitlines() if part.strip())


def replace_citations(text: str, citation_numbers: dict[str, int]) -> str:
    def repl(match):
        nums = [citation_numbers[key.strip()] for key in match.group(1).split(",")]
        return "[" + ", ".join(str(n) for n in nums) + "]"

    return re.sub(r"\\cite\{([^}]+)\}", repl, text)


def plain_markdown(text: str, citation_numbers: dict[str, int]) -> str:
    text = replace_citations(text, citation_numbers)
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"\$([^$]+)\$", lambda m: latex_math_to_text(m.group(1)), text)
    return text.strip()


INLINE = re.compile(r"(\*\*.*?\*\*|`.*?`|\$.*?\$)")


def omml_element(latex: str):
    """Convert LaTeX to a native Word Office Math (OMML) element."""
    mathml = latex_to_mathml(latex.strip())
    omml = mathml_to_omml(mathml)
    # mathml2omml 0.0.2 closes groupChrPr with the parent tag for overbars.
    omml = re.sub(r"(<m:groupChrPr>.*?)</m:groupChr>(?=<m:e>)",
                  r"\1</m:groupChrPr>", omml, flags=re.S)
    omml = omml.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
    return parse_xml(omml)


def append_rich_segment(paragraph, text: str, *, size, bold=False):
    """Append plain/code/math runs while preserving their document order."""
    for token in re.split(r"(`.*?`|\$.*?\$)", text):
        if not token:
            continue
        if token.startswith("$") and token.endswith("$"):
            paragraph._p.append(omml_element(token[1:-1]))
            continue
        latin = "Consolas" if token.startswith("`") and token.endswith("`") else FONT_LATIN
        if latin == "Consolas":
            token = token[1:-1]
        run = paragraph.add_run(token)
        set_run_font(run, cn=FONT_CN, latin=latin, size=size, bold=bold)


def add_inline(paragraph, text: str, citation_numbers: dict[str, int], *,
               size=Pt(10.5), default_bold=False):
    text = replace_citations(text, citation_numbers)
    text = text.replace(r"\%", "%")
    for token in re.split(r"(\*\*.*?\*\*)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            append_rich_segment(paragraph, token[2:-2], size=size, bold=True)
        else:
            append_rich_segment(paragraph, token, size=size, bold=default_bold)


def style_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line=Pt(21), line_spacing=1.35,
                    before=Pt(0), after=Pt(3)):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = first_line
    fmt.line_spacing = line_spacing
    fmt.space_before = before
    fmt.space_after = after
    fmt.widow_control = True


def add_body(doc, text, citation_numbers, *, indent=True, bold=False):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, first_line=Pt(21) if indent else Pt(0))
    add_inline(paragraph, text, citation_numbers, default_bold=bold)
    return paragraph


def add_heading(doc, text: str, level: int, citation_numbers):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {level}"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.keep_with_next = True
    add_inline(paragraph, text, citation_numbers,
               size={1: Pt(14), 2: Pt(12), 3: Pt(10.5)}[level], default_bold=True)
    return paragraph


def add_equation(doc, latex: str):
    if r"\begin{aligned}" in latex:
        body = latex.replace(r"\begin{aligned}", "").replace(r"\end{aligned}", "")
        lines = [part.strip().replace("&", "") for part in re.split(r"\\\\", body)
                 if part.strip()]
    else:
        lines = [latex.strip()]
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                        first_line=Pt(0), line_spacing=1.0,
                        before=Pt(3) if index == 0 else Pt(0),
                        after=Pt(5) if index == len(lines) - 1 else Pt(0))
        paragraph._p.append(omml_element(line))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, *, top=None, bottom=None):
    """Apply strict three-line-table borders to one cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)
    settings = {
        "top": top,
        "left": None,
        "bottom": bottom,
        "right": None,
        "insideH": None,
        "insideV": None,
    }
    for edge, width in settings.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single" if width else "nil")
        if width:
            element.set(qn("w:sz"), str(width))
            element.set(qn("w:color"), "000000")
            element.set(qn("w:space"), "0")
        borders.append(element)


def add_table(doc, raw_rows: list[list[str]], citation_numbers):
    rows = [raw_rows[0]] + raw_rows[2:]
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.allow_autofit = True

    for row_index, row in enumerate(rows):
        group_row = row_index > 0 and bool(row) and row[0].strip().startswith("**")
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)

        for col_index in range(cols):
            raw = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT if col_index == 0
                                   else WD_ALIGN_PARAGRAPH.CENTER)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, raw, citation_numbers,
                       size=Pt(7.5 if cols >= 8 else 8.5),
                       default_bold=(row_index == 0 or group_row))
            set_cell_borders(
                cell,
                top=12 if row_index == 0 else None,
                bottom=8 if row_index == 0 else (12 if row_index == len(rows) - 1 else None),
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc, text: str, citation_numbers, *, table=False):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.1,
                    before=Pt(5 if table else 2), after=Pt(3))
    add_inline(paragraph, text, citation_numbers, size=Pt(9), default_bold=False)
    paragraph.paragraph_format.keep_with_next = table
    return paragraph


def add_figure(doc, path: Path):
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        width_px, height_px = image.size
    width_cm = 15.5
    height_cm = width_cm * height_px / width_px
    if height_cm > 18.0:
        height_cm = 18.0
        width_cm = height_cm * width_px / height_px
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.0, before=Pt(4), after=Pt(2))
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    paragraph.paragraph_format.keep_with_next = True


def parse_bibliography() -> tuple[dict[str, int], list[tuple[str, str]]]:
    text = BBL.read_text(encoding="utf-8")
    items = re.findall(r"\\bibitem\{([^}]+)\}(.*?)(?=\\bibitem\{|\\end\{thebibliography\})",
                       text, flags=re.S)

    accents = {
        r'{\"u}': "ü", r"{\'e}": "é", r"{\'a}": "á", r"{\'i}": "í",
        r"{\'o}": "ó", r"{\'u}": "ú", r"{\~n}": "ñ",
    }

    def clean(value: str) -> str:
        value = value.replace("\n", " ")
        for old, new in accents.items():
            value = value.replace(old, new)
        value = re.sub(r"\\hskip\s+1em\s+plus\s+0\.5em\s+minus\s+0\.4em\\relax", " ", value)
        value = re.sub(r"\\(?:emph|textit|textbf|bibinfo)\{(?:[^{}]*)\}\{([^{}]*)\}", r"\1", value)
        value = re.sub(r"\\(?:emph|textit|textbf)\{([^{}]*)\}", r"\1", value)
        value = value.replace(r"\,", ",").replace("~", " ")
        value = value.replace("``", "“").replace("''", "”")
        value = value.replace("--", "–")
        value = value.replace("{", "").replace("}", "")
        value = re.sub(r"\\[A-Za-z]+", "", value)
        value = re.sub(r"\s+", " ", value).strip()
        return value

    bibliography = [(key, clean(value)) for key, value in items]
    return {key: i for i, (key, _) in enumerate(bibliography, 1)}, bibliography


def add_reference(doc, number: int, text: str):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, first_line=Pt(-18), line_spacing=1.15,
                    before=Pt(0), after=Pt(2))
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run(f"[{number}] {text}")
    set_run_font(run, cn=FONT_CN, latin=FONT_LATIN, size=Pt(9))


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_run_font(run, size=Pt(9))


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    for level, size, before, after in ((1, 14, 12, 6), (2, 12, 9, 4), (3, 10.5, 6, 3)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_HEADING
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.title = "RS-Token：面向信道鲁棒遥感通信的层次化 RemoteCLIP 蒸馏 Token"
    doc.core_properties.author = "Baohui Zhang; Jianwei Tai"
    doc.core_properties.subject = "中文论文稿 v0.6"
    doc.core_properties.keywords = "语义通信; 遥感; RVQ; RemoteCLIP; 信道鲁棒性"


def add_front_matter(doc: Document, lines: list[str], citation_numbers):
    title = lines[0][2:].strip()
    author_line = lines[2].removeprefix("作者：").strip()
    emails = re.findall(r"[\w.+'-]+@[\w.-]+", author_line)

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.25, before=Pt(0), after=Pt(10))
    run = paragraph.add_run(title)
    set_run_font(run, cn=FONT_HEADING, latin=FONT_LATIN, size=Pt(18), bold=True)

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.2, before=Pt(0), after=Pt(3))
    add_inline(paragraph, "Baohui Zhang，Jianwei Tai*", citation_numbers, size=Pt(10.5))

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.2, before=Pt(0), after=Pt(2))
    run = paragraph.add_run("安徽大学互联网学院，合肥 230601")
    set_run_font(run, size=Pt(10.5))

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line=Pt(0), line_spacing=1.2, before=Pt(0), after=Pt(12))
    contact = f"电子邮箱：{emails[0]}；* 通讯作者：{emails[1]}" if len(emails) >= 2 else author_line
    add_inline(paragraph, contact, citation_numbers, size=Pt(9.5))


def ensure_figure_four():
    path = FIG_DIR / "fig_v06_snr_sweep.png"
    if path.exists():
        return
    subprocess.run([sys.executable, str(HERE / "build_fig4_snr_sweep.py")],
                   check=True, cwd=ROOT)


def build() -> Path:
    ensure_figure_four()
    citation_numbers, bibliography = parse_bibliography()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_front_matter(doc, lines, citation_numbers)

    # Front matter metadata occupies lines 1-8. Start at the abstract heading.
    index = next(i for i, line in enumerate(lines) if line.strip() == "## 摘要")
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line or line == "---":
            index += 1
            continue
        if line == "## 参考文献":
            add_heading(doc, "参考文献", 1, citation_numbers)
            for number, (_, reference) in enumerate(bibliography, 1):
                add_reference(doc, number, reference)
            break
        if line.startswith("\\bibliography") or line.startswith("\\bibliographystyle"):
            index += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1)) - 1
            text = heading.group(2)
            if text == "摘要":
                paragraph = doc.add_paragraph()
                style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER,
                                first_line=Pt(0), before=Pt(0), after=Pt(5))
                run = paragraph.add_run("摘要")
                set_run_font(run, cn=FONT_HEADING, latin=FONT_HEADING,
                             size=Pt(12), bold=True)
            else:
                add_heading(doc, text, level, citation_numbers)
            index += 1
            continue

        if line.startswith("$$"):
            if line.endswith("$$") and len(line) > 4:
                equation = line[2:-2]
                index += 1
            else:
                pieces = []
                index += 1
                while index < len(lines) and lines[index].strip() != "$$":
                    pieces.append(lines[index])
                    index += 1
                index += 1
                equation = "\n".join(pieces)
            add_equation(doc, equation)
            continue

        figure = re.match(r"^>\s*\*\*图\s*(\d+)\*\*（`([^`]+)`）：(.+)$", line)
        if figure:
            number, filename, caption = figure.groups()
            if filename.lower().endswith(".pdf"):
                filename = str(Path(filename).with_suffix(".png"))
            add_figure(doc, FIG_DIR / filename)
            add_caption(doc, f"图 {number}  {caption}", citation_numbers)
            index += 1
            continue
        if line.startswith(">"):
            # Editorial synchronization note from the Markdown source.
            index += 1
            continue

        table_caption = re.match(r"^\*\*表\s*(\d+)\*\*：(.+)$", line)
        if table_caption:
            add_caption(doc, f"表 {table_caption.group(1)}  {table_caption.group(2)}",
                        citation_numbers, table=True)
            index += 1
            continue

        if line.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                table_rows.append(cells)
                index += 1
            add_table(doc, table_rows, citation_numbers)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph()
            style_paragraph(paragraph, first_line=Pt(0), before=Pt(0), after=Pt(2))
            paragraph.paragraph_format.left_indent = Pt(21)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            run = paragraph.add_run("• ")
            set_run_font(run, size=Pt(10.5))
            add_inline(paragraph, line[2:], citation_numbers)
            index += 1
            continue

        keywords = line.startswith("**关键词：**")
        paragraph = add_body(doc, line, citation_numbers, indent=not keywords)
        if keywords:
            paragraph.paragraph_format.space_after = Pt(8)
        index += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
